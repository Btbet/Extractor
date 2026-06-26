from fastapi import FastAPI, UploadFile, File
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from typing import List
from extract import extract_cv_data
from pypdf import PdfReader
from io import BytesIO
import json
import os
import hashlib
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib import styles
from openpyxl import Workbook
from supabase import create_client
import os
import re
from skills_db import TECHNICAL_SKILLS
from skills_db import TECHNICAL_SKILLS, SKILL_SYNONYMS

SUPABASE_URL = os.getenv("SUPABASE_URL")
SUPABASE_KEY = os.getenv("SUPABASE_KEY")

supabase = create_client(
    SUPABASE_URL,
    SUPABASE_KEY
)

from matcher import (
    extract_job_skills,
    calculate_match_score,
    generate_summary
)



def detect_skills(text):

    found = []

    text = text.lower()

    for alias, official in SKILL_SYNONYMS.items():
        text = text.replace(alias, official.lower())

    for skill in TECHNICAL_SKILLS:

        pattern = r"\b" + re.escape(skill.lower()) + r"\b"

        if re.search(pattern, text):
            found.append(skill)

    return sorted(list(set(found)))


def extract_skills_section(text: str):
    """
    Extract skills directly from the Skills section of the CV.
    Falls back to detect_skills() if no section is found.
    """

    lines = text.splitlines()

    skills = []
    collecting = False

    stop_headers = [
        "experience",
        "education",
        "project",
        "projects",
        "employment",
        "certification",
        "certifications",
        "reference",
        "references",
        "summary",
        "profile",
        "objective",
        "languages",
        "interests"
    ]

    for line in lines:

        current = line.strip()

        if not current:
            continue

        lower = current.lower()

        # Start collecting after Skills header
        if lower in [
            "skills",
            "technical skills",
            "core skills",
            "key skills",
            "competencies"
        ]:
            collecting = True
            continue

        if collecting:

            # Stop when another section begins
            if any(
                lower.startswith(header)
                for header in stop_headers
            ):
                break

            # Split commas, bullets, pipes, semicolons
            parts = re.split(
                r"[•,\|\;/]+",
                current
            )

            for part in parts:

                skill = part.strip()

                if len(skill) >= 2:
                    skills.append(skill)

    # Clean duplicates
    cleaned = []

    seen = set()

    for skill in skills:

        key = skill.lower()

        if key not in seen:
            cleaned.append(skill)
            seen.add(key)

    # Fall back to keyword detector
    if not cleaned:
        cleaned = detect_skills(text)

    return cleaned

def clean_skills(skills):

    cleaned = []

    for skill in skills:

        skill = skill.strip()

        if skill and skill not in cleaned:

            cleaned.append(skill)

    return sorted(cleaned)

app = FastAPI()
candidates_db = []

app.mount(
    "/static",
    StaticFiles(directory="."),
    name="static"
)
from fastapi.middleware.cors import CORSMiddleware

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse

app.mount("/static", StaticFiles(directory="."), name="static")

async def extract_text_from_file(file):

    content = await file.read()

    filename = file.filename.lower()

    # PDF
    if filename.endswith(".pdf"):

        reader = PdfReader(BytesIO(content))

        text = ""

        for page in reader.pages:

            extracted = page.extract_text()

            if extracted:
                text += extracted + "\n"

        return text

    # DOCX
    elif filename.endswith(".docx"):

        doc = Document(BytesIO(content))

        return "\n".join(
            p.text.strip()
            for p in doc.paragraphs
            if p.text.strip()
        )

    # TXT
    elif filename.endswith(".txt"):

        return content.decode(
            "utf-8",
            errors="ignore"
        )

    # DOC (old Microsoft Word format)
    elif filename.endswith(".doc"):

        raise Exception(
            "DOC files are not supported. Please save as DOCX or PDF."
        )

    # Unsupported file
    else:

        raise Exception(
            f"Unsupported file type: {filename}"
        )


@app.get("/")
async def dashboard():
    return FileResponse("index.html")


DB_FILE = "candidates.json"

if not os.path.exists(DB_FILE):
    with open(DB_FILE, "w") as f:
        json.dump([], f)

@app.get("/health")
async def health():
    return {
        "status": "ok"
    }
@app.post("/extract-cv")
async def extract_cv(
    file: UploadFile = File(...)
):
    try:

        # ------------------------------------------
        # Read text
        # ------------------------------------------

        text = await extract_text_from_file(file)

        # ------------------------------------------
        # Count every upload
        # ------------------------------------------

        stats = (
            supabase.table("stats")
            .select("total_uploads")
            .eq("id", 1)
            .execute()
        )

        if stats.data:

            current_uploads = stats.data[0]["total_uploads"]

            supabase.table("stats").update({
                "total_uploads": current_uploads + 1
            }).eq(
                "id",
                1
            ).execute()

        else:

            supabase.table("stats").insert({
                "id": 1,
                "total_uploads": 1
            }).execute()

        # ------------------------------------------
        # Validate document
        # ------------------------------------------

        text_lower = text.lower()

        checks = [

            "experience" in text_lower,

            "education" in text_lower,

            "skill" in text_lower,

            "@" in text

        ]

        if sum(checks) < 3:

            return {

                "status": "rejected",

                "message": "Invalid CV"

            }

        # ------------------------------------------
        # Extract candidate
        # ------------------------------------------

        candidate = extract_cv_data(text)

        # ------------------------------------------
        # Always store education as an array
        # ------------------------------------------

        education = candidate.get("education", [])

        if isinstance(education, str):

            education = [education]

        elif education is None:

            education = []

        candidate["education"] = education
    # ----------------------------
# Skills
# ----------------------------

skills = extract_skills_section(text)

# If no Skills section was found, scan the whole CV
if not skills:
    skills = detect_skills(text)

candidate["skills"] = clean_skills(skills)

        candidate["summary"] = generate_summary(candidate)

        cv_hash = hashlib.sha256(
            text.encode("utf-8")
        ).hexdigest()

        candidate["cv_hash"] = cv_hash

        # ------------------------------------------
        # Check duplicate
        # ------------------------------------------

        response = (
            supabase.table("candidates")
            .select("*")
            .execute()
        )

        for c in response.data:

            if (

                (
                    c.get("email")
                    and candidate.get("email")
                    and c["email"].lower()
                    == candidate["email"].lower()
                )

                or

                (
                    c.get("cv_hash")
                    and candidate.get("cv_hash")
                    and c["cv_hash"]
                    == candidate["cv_hash"]
                )

            ):

                return {

                    "status": "duplicate",

                    "message": "Duplicate CV"

                }

        # ------------------------------------------
        # Save candidate
        # ------------------------------------------

        supabase.table("candidates").insert({

            "name": candidate.get("name"),

            "email": candidate.get("email"),

            "phone": candidate.get("phone"),

            "skills": candidate.get("skills"),

            "education": candidate.get("education"),

            "years_experience": candidate.get(
                "years_experience"
            ),

            "score": candidate.get("score"),

            "summary": candidate.get("summary"),

            "cv_hash": candidate.get("cv_hash")

        }).execute()

        # ------------------------------------------
        # Success
        # ------------------------------------------

        return {

            "status": "success",

            "message": "CV Uploaded Successfully",

            "name": candidate.get("name"),

            "email": candidate.get("email"),

            "candidate": candidate

        }

    except Exception as e:

        print(f"Upload error: {e}")

        return {

            "status": "failed"

        }

@app.post("/upload-multiple")
async def upload_multiple(
    files: List[UploadFile] = File(...)
):

    uploaded = []
    duplicates = []
    rejected = []
    failed = []

    for file in files:

        try:

            # ------------------------------------------
            # Read file
            # ------------------------------------------

            text = await extract_text_from_file(file)

            # ------------------------------------------
            # Count every upload
            # ------------------------------------------

            stats = (
                supabase.table("stats")
                .select("total_uploads")
                .eq("id", 1)
                .single()
                .execute()
            )

            current_uploads = stats.data["total_uploads"]

            supabase.table("stats").update({
                "total_uploads": current_uploads + 1
            }).eq(
                "id",
                1
            ).execute()

            # ------------------------------------------
            # Validate document
            # ------------------------------------------

            text_lower = text.lower()

            checks = [

                "experience" in text_lower,

                "education" in text_lower,

                "skill" in text_lower,

                "@" in text

            ]

            if sum(checks) < 3:

                rejected.append(file.filename)

                continue

            # ------------------------------------------
            # Extract candidate
            # ------------------------------------------

            candidate = extract_cv_data(text)

            candidate["skills"] = detect_skills(text)

            candidate["skills"] = clean_skills(
                candidate.get("skills", [])
            )

            candidate["summary"] = generate_summary(candidate)

            cv_hash = hashlib.sha256(
                text.encode("utf-8")
            ).hexdigest()

            candidate["cv_hash"] = cv_hash

            # ------------------------------------------
            # Check duplicate
            # ------------------------------------------

            duplicate = False

            response = (
                supabase.table("candidates")
                .select("*")
                .execute()
            )

            for c in response.data:

                if (

                    (
                        c.get("email")
                        and candidate.get("email")
                        and c["email"].lower()
                        == candidate["email"].lower()
                    )

                    or

                    (
                        c.get("cv_hash")
                        and candidate.get("cv_hash")
                        and c["cv_hash"]
                        == candidate["cv_hash"]
                    )

                ):

                    duplicate = True
                    break

            if duplicate:

                duplicates.append(file.filename)

                continue

            # ------------------------------------------
            # Save candidate
            # ------------------------------------------

            supabase.table("candidates").insert({

                "name": candidate.get("name"),

                "email": candidate.get("email"),

                "phone": candidate.get("phone"),

                "skills": candidate.get("skills"),

                "education": candidate.get("education"),

                "years_experience": candidate.get(
                    "years_experience"
                ),

                "score": candidate.get("score"),

                "summary": candidate.get("summary"),

                "cv_hash": candidate.get("cv_hash")

            }).execute()

            uploaded.append(file.filename)

        except Exception as e:

            print(f"Error in {file.filename}: {e}")

            failed.append(file.filename)

    return {

        "status": "completed",

        "message": "Upload Complete",

        "uploaded": uploaded,

        "duplicates": duplicates,

        "rejected": rejected,

        "failed": failed,

        "saved_count": len(uploaded),

        "duplicate_count": len(duplicates),

        "rejected_count": len(rejected),

        "failed_count": len(failed),

        "total_processed": len(files)

    }

@app.get("/candidates")
def get_candidates(page: int = 1, limit: int = 10):

    start = (page - 1) * limit
    end = start + limit - 1

    response = (
        supabase.table("candidates")
        .select("*", count="exact")
        .range(start, end)
        .execute()
    )

    candidates = response.data
    total = response.count

    output = []

    for i, candidate in enumerate(candidates, start=start + 1):
        candidate["number"] = i
        output.append(candidate)

    return {
        "total": total,
        "page": page,
        "limit": limit,
        "pages": (total + limit - 1) // limit,
        "candidates": output
    }

@app.get("/search")
def search(query: str):

    q = query.strip().lower()

    # Search by name in Supabase
    response = (
        supabase.table("candidates")
        .select("*")
        .ilike("name", f"%{q}%")
        .execute()
    )

    results = response.data or []

    # Search skills in Python
    response = (
        supabase.table("candidates")
        .select("*")
        .execute()
    )

    for candidate in response.data:

        if candidate in results:
            continue

        skills = candidate.get("skills", [])

        if isinstance(skills, str):
            skills = [skills]

        if any(q in skill.lower() for skill in skills):
            results.append(candidate)

    for i, candidate in enumerate(results, start=1):
        candidate["number"] = i

    return results
@app.get("/download-match-doc")
def download_match_doc():

    if not os.path.exists(
        "matched_results.json"
    ):

        return {
            "error":"Run Match Job first"
        }

    with open(
        "matched_results.json",
        "r"
    ) as f:

        candidates=json.load(f)

    doc=Document()

    doc.add_heading(
        "Matched Candidates Report",
        level=1
    )

    for index, c in enumerate(
        candidates,
        start=1
    ):

        doc.add_heading(
            f"{index}. {c.get('name','N/A')}",
            level=2
        )

        doc.add_paragraph(
            f"Match Score: {c.get('job_match_score',0)}%"
        )

        doc.add_paragraph(
            "Matched Skills: " +
            ", ".join(
                c.get(
                    "matched_skills",
                    []
                )
            )
        )

        doc.add_paragraph(
            "AI Summary: " +
            str(
                c.get(
                    "summary",
                    ""
                )
            )
        )

    filename="results.docx"

    doc.save(
        filename
    )

    return FileResponse(
        filename,
        filename="results.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
@app.get("/download-match-pdf")
def download_match_pdf():

    if not os.path.exists(
        "matched_results.json"
    ):

        return {
            "error":"Run Match Job first"
        }

    with open(
        "matched_results.json",
        "r"
    ) as f:

        candidates=json.load(f)

    pdf_file="results.pdf"

    doc=SimpleDocTemplate(
        pdf_file
    )

    stylesheet=styles.getSampleStyleSheet()

    normal=stylesheet["Normal"]

    title=stylesheet["Title"]

    story=[]

    story.append(
        Paragraph(
            "Matched Candidates Report",
            title
        )
    )

    story.append(
        Spacer(1,20)
    )

    for index, c in enumerate(
        candidates,
        start=1
    ):

        story.append(
            Paragraph(
                f"{index}. Name: {c.get('name', 'N/A')}",
                normal
            )
        )

        story.append(
            Paragraph(
                f"Match Score: {c.get('job_match_score',0)}%",
                normal
            )
        )

        story.append(
            Paragraph(
                "Matched Skills: " +
                ", ".join(
                    c.get(
                        "matched_skills",
                        []
                    )
                ),
                normal
            )
        )

        story.append(
            Paragraph(
                "AI Summary: " +
                str(
                    c.get(
                        "summary",
                        ""
                    )
                ),
                normal
            )
        )

        story.append(
            Spacer(1,15)
        )

    doc.build(
        story
    )

    return FileResponse(
        pdf_file,
        filename="results.pdf",
        media_type="application/pdf"
    )

@app.get("/export-csv")
def export_csv():

    # Read candidates from Supabase
    response = supabase.table("candidates").select("*").execute()

    candidates = response.data

    wb = Workbook()

    ws = wb.active

    ws.title = "candidates"

    ws.append([

        "Name",
        "Email",
        "Phone",
        "Skills",
        "Education",
        "Experience",
        "Score"

    ])

    for c in candidates:

        ws.append([

            c.get(
                "name",
                ""
            ),

            c.get(
                "email",
                ""
            ),

            c.get(
                "phone",
                ""
            ),

            " | ".join(
                c.get(
                    "skills",
                    []
                )
            ),

            " | ".join(
                c.get(
                    "education",
                    []
                )
            ),

            str(
                c.get(
                    "years_experience",
                    0
                )
            ) + " years",

            c.get(
                "score",
                0
            )

        ])

    ws.column_dimensions["A"].width = 25
    ws.column_dimensions["B"].width = 35
    ws.column_dimensions["C"].width = 20
    ws.column_dimensions["D"].width = 50
    ws.column_dimensions["E"].width = 35
    ws.column_dimensions["F"].width = 15
    ws.column_dimensions["G"].width = 10

    wb.save(
        "candidates.xlsx"
    )

    return FileResponse(

        "candidates.xlsx",

        filename="candidates.xlsx",

        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"

    )

@app.get("/export-pdf")
def export_pdf():

    # Read candidates from Supabase
    response = supabase.table("candidates").select("*").execute()

    candidates = response.data

    pdf_file = "candidates.pdf"

    doc = SimpleDocTemplate(
        pdf_file
    )

    stylesheet = styles.getSampleStyleSheet()

    normal = stylesheet["Normal"]
    title = stylesheet["Title"]

    story = []

    story.append(
        Paragraph(
            "Candidates Report",
            title
        )
    )

    story.append(
        Spacer(1,20)
    )

    for index, c in enumerate(
        candidates,
        start=1
    ):

        story.append(
            Paragraph(
                f"{index}. {c.get('name', 'N/A')}",
                normal
            )
        )

        story.append(
            Paragraph(
                f"Email: {c.get('email','')}",
                normal
            )
        )

        story.append(
            Paragraph(
                "Skills: " +
                ", ".join(
                    c.get(
                        "skills",
                        []
                    )
                ),
                normal
            )
        )

        story.append(
            Spacer(1,15)
        )

    doc.build(story)

    return FileResponse(
        pdf_file,
        filename="candidates.pdf",
        media_type="application/pdf"
    )

@app.post("/match-job")
def match_job(description: str):

    # Read candidates from Supabase
    response = supabase.table("candidates").select("*").execute()
    candidates = response.data

    required_skills = extract_job_skills(description)

    results = []

    for candidate in candidates:

        result = calculate_match_score(

            candidate.get(
                "skills",
                []
            ),

            required_skills

        )

        candidate_copy = candidate.copy()

        candidate_copy[
            "job_match_score"
        ] = result["score"]

        candidate_copy[
            "matched_skills"
        ] = result["matched"]

        candidate_copy[
            "missing_skills"
        ] = result["missing"]

        candidate_copy[
            "related_skills"
        ] = result["related"]

        candidate_copy[
            "comments"
        ] = result["comments"]

        candidate_copy[
            "summary"
        ] = generate_summary(
            candidate
        )

        results.append(
            candidate_copy
        )

    results = sorted(

        results,

        key=lambda x: x["job_match_score"],

        reverse=True

    )

    # Save matched results locally for PDF/DOC export
    with open(
        "matched_results.json",
        "w"
    ) as f:

        json.dump(
            results,
            f,
            indent=4
        )

    return {

        "required_skills": required_skills,

        "ranked_candidates": results

    }

@app.get("/debug-candidates")
def debug_candidates():

    response = supabase.table(
        "candidates"
    ).select("*").execute()

    return response.data
    
@app.delete("/reset_session")
def reset_session():

    with open(
        DB_FILE,
        "w"
    ) as f:

        json.dump(
            [],
            f
        )

    with open(
        STATS_FILE,
        "w"
    ) as f:

        json.dump(
            {
                "total_uploads": 0
            },
            f
        )

    return {

        "message":
        "Session cleared successfully"

    }
@app.get("/total-uploads")
def total_uploads():

    # Get total uploads from stats table
    stats = (
        supabase.table("stats")
        .select("total_uploads")
        .eq("id", 1)
        .single()
        .execute()
    )

    total_uploads = 0

    if stats.data:
        total_uploads = stats.data["total_uploads"]

    # Get unique candidates
    candidates = (
        supabase.table("candidates")
        .select("*", count="exact")
        .execute()
    )

    total_candidates = candidates.count

    return {
        "total_uploads": total_uploads,
        "total_candidates": total_candidates
    }
